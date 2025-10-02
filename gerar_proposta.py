# gerar_proposta.py
from docx import Document
import io, os, uuid
from tempfile import TemporaryDirectory
from shutil import copyfile
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx.opc.constants
import subprocess, shutil  # para converter via LibreOffice no Linux/macOS

# ─── helpers PDF ─────────────────────────────────────────────────────────────
try:
    from docx2pdf import convert
    import pythoncom
    _DOCX2PDF_AVAILABLE = True
except ImportError:
    _DOCX2PDF_AVAILABLE = False

try:
    import pdfkit  # não usamos aqui, mas mantido para compatibilidade
    _PDFKIT_AVAILABLE = True
except ImportError:
    pdfkit = None
    _PDFKIT_AVAILABLE = False
# ─────────────────────────────────────────────────────────────────────────────

# Base do projeto (para resolver caminhos relativos de imagens)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def _resolve_img_path(pth: str):
    """
    Resolve um caminho absoluto para a imagem do equipamento,
    normalizando separadores e tentando algumas pastas comuns
    (CWD, BASE_DIR, BASE_DIR/static[/images]).
    """
    if not pth:
        return None

    text = str(pth).strip()
    if not text:
        return None

    # Se já for absoluto e existir, ok
    if os.path.isabs(text) and os.path.exists(text):
        return text

    # Normaliza separadores vindos de Windows e remove prefixos redundantes
    normalized = text.replace("\\", "/").lstrip("/")
    parts = [p for p in normalized.split("/") if p and p not in (".", "..")]

    trimmed = parts[:]
    if trimmed and trimmed[0].lower() == "static":
        trimmed = trimmed[1:]
    if trimmed and trimmed[0].lower() == "images":
        trimmed = trimmed[1:]

    joined = os.path.join(*parts) if parts else ""
    trimmed_joined = os.path.join(*trimmed) if trimmed else ""
    basename = trimmed[-1] if trimmed else (parts[-1] if parts else "")

    candidates = []
    seen_rel = set()
    for rel in (joined, trimmed_joined):
        if not rel or rel in seen_rel:
            continue
        seen_rel.add(rel)
        candidates.extend([
            rel,
            os.path.join(os.getcwd(), rel),
            os.path.join(BASE_DIR, rel),
        ])

    if trimmed_joined:
        candidates.extend([
            os.path.join(BASE_DIR, "static", trimmed_joined),
            os.path.join(BASE_DIR, "static", "images", trimmed_joined),
        ])

    if basename:
        candidates.append(os.path.join(BASE_DIR, "static", "images", basename))

    seen = set()
    for candidate in candidates:
        candidate_path = os.path.normpath(candidate)
        if candidate_path in seen:
            continue
        seen.add(candidate_path)
        if os.path.exists(candidate_path):
            return os.path.abspath(candidate_path)

    return None


# --------------------------------------------------------------------------- #
# Utilitários de telefone / hyperlink
# --------------------------------------------------------------------------- #
def _clean_phone(raw: str) -> str:
    """Deixa só dígitos."""
    return "".join(filter(str.isdigit, raw or ""))


def _valid_phone(digits: str) -> bool:
    """Considera válido se possuir pelo menos 12 dígitos (DDI+DDD+celular)."""
    return len(digits) >= 12


def _add_hyperlink(paragraph, url, text):
    """Insere hyperlink preservando estilo."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")  # mantém formatação corrente
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hl.append(new_run)
    paragraph._p.append(hl)
    return hl


def _linkify_phone(doc: Document, raw_phone: str, digits: str):
    """Procura texto 'Telefone: <raw_phone>' e transforma em link WhatsApp."""
    wa_url = f"https://wa.me/{digits}"
    for para in doc.paragraphs:
        if raw_phone in para.text:
            # divide texto mantendo prefixo/sufixo
            parts = para.text.split(raw_phone)
            # limpa runs
            for r in para.runs:
                r.text = ""
            # reconstrói com hyperlink no meio
            if parts[0]:
                para.add_run(parts[0])
            _add_hyperlink(para, wa_url, raw_phone)
            if len(parts) > 1 and parts[1]:
                para.add_run(parts[1])
            break


# --------------------------------------------------------------------------- #
# Substituição de {{ campos }}
# --------------------------------------------------------------------------- #
def _substituir_campos(doc, mapa):
    for p in doc.paragraphs:
        _replace(p, mapa)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace(p, mapa)


def _replace(paragraph, mapa):
    txt = paragraph.text
    for k, v in mapa.items():
        token = f"{{{{ {k} }}}}"
        if token in txt:
            txt = txt.replace(token, str(v))
    if txt != paragraph.text:
        for r in paragraph.runs:
            r.text = ""
        (paragraph.runs[0] if paragraph.runs else paragraph.add_run()).text = txt


# --------------------------------------------------------------------------- #
# Tabela de equipamentos
# --------------------------------------------------------------------------- #
def _inserir_tabela_equipamentos(doc, equipamentos):
    """
    Cria tabela, adiciona coluna de desconto só se houver, e insere após a
    âncora 'INVESTIMENTO' (tolerante: aceita 'INVESTIMENTO:' e
    'INVESTIMENTO (AQUISIÇÃO):', case-insensitive).
    """
    # há algum item com desconto?
    has_discount = any((getattr(eq, "discount_percent", 0) or 0) != 0 for eq in equipamentos)

    # procura âncora
    texto_alvo = ("INVESTIMENTO:", "INVESTIMENTO (AQUISIÇÃO):")
    anchor_i = next(
        (i for i, p in enumerate(doc.paragraphs)
         if any(t in (p.text or "").upper() for t in texto_alvo)),
        None
    )

    # cria a tabela (por padrão, vai para o fim)
    cols = 6 if has_discount else 5
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Descrição"
    hdr[1].text = "Imagem"
    hdr[2].text = "Quantidade"
    hdr[3].text = "Preço Unitário"
    if has_discount:
        hdr[4].text = "Preço c/ desconto"
        hdr[5].text = "Total"
        cent_cols = (1, 2, 3, 4, 5)
    else:
        hdr[4].text = "Total"
        cent_cols = (1, 2, 3, 4)

    # centraliza cabeçalho
    for c in hdr:
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in c.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # linhas
    for eq in equipamentos:
        pct   = float(getattr(eq, "discount_percent", 0) or 0.0)
        cheio = float(getattr(eq, "unit_price", 0) or 0.0)
        qtd   = int(getattr(eq, "quantity", 1) or 1)
        desc  = cheio * (1 - pct/100.0)
        sub   = desc * qtd

        row = table.add_row().cells
        row[0].text = (getattr(eq, "description", None) or getattr(eq, "name", "") or "")

        # imagem (resolve caminho absoluto de forma robusta)
        img_path = _resolve_img_path(getattr(eq, "illustration_path", None))
        if img_path:
            run = row[1].paragraphs[0].add_run()
            # 160 px ~ 1.67" @96dpi (a imagem já vem cortada para 160x180 pelo upload)
            run.add_picture(img_path, width=Inches(1.67))
        else:
            row[1].text = "—"

        row[2].text = str(qtd)
        row[3].text = _fmt(cheio)

        if has_discount:
            row[4].text = _fmt(desc) if pct else ""
            row[5].text = _fmt(sub)
        else:
            row[4].text = _fmt(sub)

        # centraliza colunas numéricas
        for i in cent_cols:
            row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in row[i].paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # insere a tabela logo após a âncora, se ela existir
    if anchor_i is not None and 0 <= anchor_i < len(doc.paragraphs):
        doc.paragraphs[anchor_i]._element.addnext(table._element)
    # se não achou, a tabela já ficou no fim do documento


def _fmt(num):
    return f"R$ {num:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


# --------------------------------------------------------------------------- #
# Função principal
# --------------------------------------------------------------------------- #
def gerar_proposta_docx(
    proposta,
    equipamentos,
    formato: str = "docx",
    *,
    nome_colaborador: str = "",
    proposta_cod: str = "",
    email_colaborador: str = "",
):
    template_path = "docs_templates/proposta_template.docx"
    if not os.path.exists(template_path):
        raise FileNotFoundError("Template DOCX não encontrado.")

    # --- valida telefone ---------------------------------------------------
    tel_raw   = proposta.telefone or ""
    tel_clean = _clean_phone(tel_raw)
    if tel_raw and not _valid_phone(tel_clean):
        raise ValueError(
            "Telefone inválido. Informe DDI+DDD+número, "
            "por exemplo: +55 11 912345678"
        )

    with TemporaryDirectory() as tmp:
        tmp_docx = os.path.join(tmp, f"{uuid.uuid4()}.docx")
        tmp_pdf  = os.path.join(tmp, f"{uuid.uuid4()}.pdf")
        copyfile(template_path, tmp_docx)
        doc = Document(tmp_docx)

        dados_topo = (
            f"{proposta.company} / {proposta.cnpj} / {proposta.client_name} / "
            f"{proposta.data_criacao.strftime('%d/%m/%Y') if proposta.data_criacao else ''}\n"
            f"Telefone: {tel_raw}  E-mail: {proposta.email}"
        )
        condicoes = (
            "CONDIÇÕES COMERCIAIS:\n"
            f". Condições de Pagamento (Equipamento): {proposta.pagamento or ''}\n"
            f". Prazo de entrega: {proposta.prazo_entrega or ''}\n"
            f". Frete: {proposta.frete or ''}\n"
            f". Validade da Proposta: {proposta.validade or ''}\n"
            f". Garantia do Equipamento: {proposta.garantia or ''}\n"
            f". Garantia do Sistema: {proposta.garantia_sistema or ''}"
        )
        mapa = {
            "empresa": proposta.company,
            "cnpj": proposta.cnpj,
            "cliente": proposta.client_name,
            "email": proposta.email,
            "telefone": tel_raw,
            "numero": tel_raw,
            "pagamento": proposta.pagamento,
            "prazo_entrega": proposta.prazo_entrega,
            "frete": proposta.frete,
            "validade": proposta.validade,
            "garantia": proposta.garantia,
            "garantia_sistema": proposta.garantia_sistema,
            "proposta_cod": proposta_cod,
            "condicoes_comerciais": condicoes,
            "nome_colaborador": nome_colaborador,
            "email_colaborador": email_colaborador,
            "data": proposta.data_criacao.strftime("%d/%m/%Y") if proposta.data_criacao else "",
            "dados_topo": dados_topo,
        }

        _substituir_campos(doc, mapa)
        _inserir_tabela_equipamentos(doc, equipamentos)

        # transforma telefone em link WhatsApp, se válido
        if _valid_phone(tel_clean):
            _linkify_phone(doc, tel_raw, tel_clean)

        doc.save(tmp_docx)

        if formato.lower() == "pdf":
            # Windows: usa Word (docx2pdf)
            if os.name == "nt" and _DOCX2PDF_AVAILABLE:
                pythoncom.CoInitialize()
                convert(tmp_docx, tmp_pdf)
                pythoncom.CoUninitialize()
                with open(tmp_pdf, "rb") as f:
                    return io.BytesIO(f.read())

            # Linux/macOS: usa LibreOffice headless (wkhtmltopdf não lê DOCX)
            soffice = shutil.which("soffice") or shutil.which("libreoffice")
            if not soffice:
                raise RuntimeError(
                    "Conversão para PDF indisponível. Instale o LibreOffice:\n"
                    "sudo apt-get install -y libreoffice-core libreoffice-writer"
                )
            cmd = [
                soffice, "--headless",
                "--convert-to", "pdf:writer_pdf_Export",
                "--outdir", tmp, tmp_docx,
            ]
            proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            if proc.returncode != 0:
                raise RuntimeError(
                    "Falha ao converter DOCX → PDF via LibreOffice.\n"
                    f"stdout:\n{proc.stdout.decode(errors='ignore')}\n\n"
                    f"stderr:\n{proc.stderr.decode(errors='ignore')}"
                )

            produced_pdf = os.path.splitext(tmp_docx)[0] + ".pdf"
            with open(produced_pdf, "rb") as f:
                return io.BytesIO(f.read())

        # Se não for PDF, retorna o DOCX
        with open(tmp_docx, "rb") as f:
            return io.BytesIO(f.read())
